import os
import re
import docx # Cần cài đặt: pip install python-docx
import pandas as pd # Cần cài đặt: pip install pandas
import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
def extract_project_info(doc_path):
    try:
        doc = docx.Document(doc_path)
        full_text = '\n'.join([para.text for para in doc.paragraphs])
        
        
        # Lấy danh sách các đoạn văn đã được làm sạch để dùng trong vòng lặp
        paragrahps = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        # Dictionary để lưu trữ dữ liệu trích xuất được
        data = {
            'English': '',
            'Context_Objectives': '',
            'Technology/algorithm': '',
            'Summarize_the_contents':'',
            'Expected features': '',          
        }
        # 2. Dùng REGEX để lấy các trường đơn giản
        english_match = re.search(r'English:\s*(.*?)\n', full_text, re.IGNORECASE)
        if english_match:
            data['English'] = english_match.group(1).strip()
            if not data['English'].endswith('.'):
                data['English'] += '.'
            data['English'] += ' ' 
        #3. Dùng VÒNG LẶP VÀ CỜ TRẠNG THÁI cho các trường phức tạp
        current_section = None
        for para_text in paragrahps:
            text_lower = para_text.strip().lower()
            is_header = False
            # Kiểm tra xem đoạn văn có chứa tiêu đề của các mục không
            if 'context:' in text_lower or 'objectives:' in text_lower:
                current_section = 'Context_Objectives'
                is_header = True
               
            elif 'technology/algorithm' in text_lower:
                current_section = 'Technology/algorithm'
                is_header = True
                
            elif 'summarize the contents to be researched and the expected outputs of the project' in text_lower:
                current_section = 'Summarize_the_contents'
                is_header = True
               
            elif 'expected features' in text_lower:
                current_section = 'Expected features'
                is_header = True
            if is_header:
                continue
            
            #nếu đã ở trong một mục, nối văn bản lại
            if current_section:
                lines = para_text.split('\n')
                processed_lines = []
                for line in lines:
                    cleaned_line = line.strip('●•*-_ ').strip()
                    
                    if cleaned_line:  # Chỉ thêm nếu dòng không rỗng
                        if not cleaned_line.endswith(('.', '?', '!',':')):
                            cleaned_line += '.'
                        processed_lines.append(cleaned_line)
                final_text_chunk = " ".join(processed_lines)
                
                if final_text_chunk:
                    if data[current_section]:
                        data[current_section] += " " + final_text_chunk
                    else:
                        data[current_section] = final_text_chunk
            
        # Loại bỏ khoảng trắng thừa ở đầu và cuối mỗi trường
        for key in data:
            data[key] = data[key].strip()
        return data
    except FileNotFoundError:
        print(f"Lỗi: Không tìm thấy file tại đường dẫn '{doc_path}'")
        return None
    except Exception as e:
        print(f"Đã xảy ra lỗi: {e}")
        return None 

# --- PHẦN CHÍNH ĐỂ XỬ LÝ THƯ MỤC VÀ XUẤT CSV ---
if __name__ == "__main__":
    # THAY ĐỔI CÁC ĐƯỜNG DẪN NÀY CHO PHÙ HỢP VỚI MÁY CỦA BẠN
    docx_folder_path = "D:/Traning AI/OK"  # Thư mục chứa các file .docx
    output_csv_file = "D:/Traning AI/extracted_projects_data2.csv"     # Tên file CSV_backup đầu ra

    all_extracted_data = [] # List để lưu dữ liệu từ tất cả các file

    print(f"Bắt đầu quét các file .docx trong thư mục: '{os.path.abspath(docx_folder_path)}'")

    # Kiểm tra xem thư mục có tồn tại không
    if not os.path.isdir(docx_folder_path):
        print(f"Lỗi: Thư mục '{docx_folder_path}' không tồn tại.")
    else:
        # Lặp qua tất cả các file trong thư mục đã chỉ định
        for filename in sorted(os.listdir(docx_folder_path)): # sorted để có thứ tự nhất quán
            if filename.lower().endswith(".docx"):
                full_file_path = os.path.join(docx_folder_path, filename)
                print(f"Đang xử lý file: {filename}...")
                
                project_info = extract_project_info(full_file_path)
                
                if project_info:
                    all_extracted_data.append(project_info)
                    print(f"  -> Đã trích xuất thành công: {filename}")
                else:
                    print(f"  -> Không thể trích xuất dữ liệu từ: {filename}")

        # Kiểm tra xem có dữ liệu nào được trích xuất không
        if all_extracted_data:
            # Chuyển list các dictionary thành một Pandas DataFrame
            df = pd.DataFrame(all_extracted_data)
            
            # Sắp xếp lại các cột nếu muốn (tùy chọn)
            # Ví dụ: df = df[['FileName', 'English', 'Context_Objectives', ...]]
            
            # Lưu DataFrame vào file CSV_backup
            try:
                df.to_csv(output_csv_file, index=False, encoding='utf-8-sig')
                print(f"\nĐã lưu thành công dữ liệu vào file: '{output_csv_file}'")
            except Exception as e:
                print(f"\nLỗi khi lưu file CSV: {e}")
        else:
            print("\nKhông có dữ liệu nào được trích xuất để lưu vào CSV.")