import docx
import csv
import re
import sys
import io


sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
def extract_project_info(doc_path):
    try:
        doc = docx.Document(doc_path)
        full_text = '\n'.join([para.text for para in doc.paragraphs])
        
        
        # Lấy danh sách các đoạn văn đã được làm sạch để dùng trong vòng lặp
        paragrahps = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        # Dictionary để lưu trữ dữ liệu trích xuất được
        data = {
            'English': '',
            'Context_Objectives': '',
            'Technology/algorithm': '',
            'Summarize_the_contents':'',
            'Expected features': '',          
        }
        # 2. Dùng REGEX để lấy các trường đơn giản
        english_match = re.search(r'English:\s*(.*?)\n', full_text, re.IGNORECASE)
        if english_match:
            data['English'] = english_match.group(1).strip()
            if not data['English'].endswith('.'):
                data['English'] += '.'
            data['English'] += ' ' 
        #3. Dùng VÒNG LẶP VÀ CỜ TRẠNG THÁI cho các trường phức tạp
        current_section = None
        for para_text in paragrahps:
            text_lower = para_text.strip().lower()
            is_header = False
            # Kiểm tra xem đoạn văn có chứa tiêu đề của các mục không
            if 'context:' in text_lower or 'objectives:' in text_lower:
                current_section = 'Context_Objectives'
                is_header = True
               
            elif 'technology/algorithm:' in text_lower:
                current_section = 'Technology/algorithm'
                is_header = True
                
            elif 'summarize the contents to be researched and the expected outputs of the project:' in text_lower:
                current_section = 'Summarize_the_contents'
                is_header = True
               
            elif 'expected features' in text_lower:
                current_section = 'Expected features'
                is_header = True
            if is_header:
                continue
            
            #nếu đã ở trong một mục, nối văn bản lại
            if current_section:
                lines = para_text.split('\n')
                processed_lines = []
                for line in lines:
                    cleaned_line = line.strip('●•- ').strip()
                    
                    if cleaned_line:  # Chỉ thêm nếu dòng không rỗng
                        if not cleaned_line.endswith(('.', '?', '!',':')):
                            cleaned_line += '.'
                        processed_lines.append(cleaned_line)
                final_text_chunk = " ".join(processed_lines)
                
                if final_text_chunk:
                    if data[current_section]:
                        data[current_section] += " " + final_text_chunk
                    else:
                        data[current_section] = final_text_chunk
            
        # Loại bỏ khoảng trắng thừa ở đầu và cuối mỗi trường
        for key in data:
            data[key] = data[key].strip()
        return data
    except FileNotFoundError:
        print(f"Lỗi: Không tìm thấy file tại đường dẫn '{doc_path}'")
        return None
    except Exception as e:
        print(f"Đã xảy ra lỗi: {e}")
        return None 
        
def write_to_csv_final(data, csv_path):
    """Ghi dữ liệu vào file CSV."""
    if not data:
        print("Không có dữ liệu để ghi.")
        return  
    # Lấy tiêu đề cột từ các khóa của dictionary
    fieldnames = list(data.keys())
    
    try:
        with open(csv_path, mode='w', newline='', encoding='utf-8-sig') as csv_file:
            writer = csv.DictWriter(csv_file, fieldnames=fieldnames)
            writer.writeheader()
            writer.writerow(data)
        print(f"Đã ghi dữ liệu thành công vào file '{csv_path}'")
    except Exception as e:
        print(f"Lỗi khi ghi file CSV: {e}")

# --- Chương trình chính ---
if __name__ == "__main__":
    # Đường dẫn đến file Word có cấu trúc đầy đủ
    # word_document_path = 'D:/SEP490/Capstone Project Registration.docx'
    word_document_path = "D:/Traning AI/ListDataCapstoneProjects/SE_19 Nhóm Lê Viết Hậu-RinNV.docx"
    
    # Đường dẫn file CSV bạn muốn tạo
    csv_output_path = 'D:/SEP490/project_data_final1.csv'
    
    # # Trích xuất và ghi dữ liệu
    project_data = extract_project_info(word_document_path)
    print("length:", len(project_data['Expected features'])/4)
    if project_data:
        write_to_csv_final(project_data, csv_output_path)   
   
   