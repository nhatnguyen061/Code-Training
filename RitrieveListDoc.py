import docx
import csv
import re
import sys
import io
import pandas as pd
import os
import random



sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

def extract_project_info(doc_path):
    try:
        doc = docx.Document(doc_path)
        full_text = '\n'.join([para.text for para in doc.paragraphs])
        
        
        # Lấy danh sách các đoạn văn đã được làm sạch để dùng trong vòng lặp
        paragrahps = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        # Dictionary để lưu trữ dữ liệu trích xuất được
        data = {
            'English': '',
            'Context_Objectives': '',
            'Technology/algorithm': '',
            'Summarize_the_contents':'',
            'Expected features': '',          
        }
        # 2. Dùng REGEX để lấy các trường đơn giản
        english_match = re.search(r'English:\s*(.*?)\n', full_text, re.IGNORECASE)
        if english_match:
            data['English'] = english_match.group(1).strip()
            if not data['English'].endswith('.'):
                data['English'] += '.'
            data['English'] += ' ' 
        #3. Dùng VÒNG LẶP VÀ CỜ TRẠNG THÁI cho các trường phức tạp
        current_section = None
        for para_text in paragrahps:
            text_lower = para_text.strip().lower()
            is_header = False
            # Kiểm tra xem đoạn văn có chứa tiêu đề của các mục không
            if 'context:' in text_lower or 'objectives:' in text_lower:
                current_section = 'Context_Objectives'
                is_header = True
               
            elif 'technology/algorithm' in text_lower:
                current_section = 'Technology/algorithm'
                is_header = True
                
            elif 'summarize the contents to be researched and the expected outputs of the project' in text_lower:
                current_section = 'Summarize_the_contents'
                is_header = True
               
            elif 'expected features' in text_lower:
                current_section = 'Expected features'
                is_header = True
            if is_header:
                continue
            
            #nếu đã ở trong một mục, nối văn bản lại
            if current_section:
                lines = para_text.split('\n')
                processed_lines = []
                for line in lines:
                    cleaned_line = line.strip('●•*-_ ').strip()
                    
                    if cleaned_line:  # Chỉ thêm nếu dòng không rỗng
                        if not cleaned_line.endswith(('.', '?', '!',':')):
                            cleaned_line += '.'
                        processed_lines.append(cleaned_line)
                final_text_chunk = " ".join(processed_lines)
                
                if final_text_chunk:
                    if data[current_section]:
                        data[current_section] += " " + final_text_chunk
                    else:
                        data[current_section] = final_text_chunk
            
        # Loại bỏ khoảng trắng thừa ở đầu và cuối mỗi trường
        for key in data:
            data[key] = data[key].strip()
        return data
    except FileNotFoundError:
        print(f"Lỗi: Không tìm thấy file tại đường dẫn '{doc_path}'")
        return None
    except Exception as e:
        print(f"Đã xảy ra lỗi: {e}")
        return None 
    

#3. Tạo ra các bộ ba
## close
# print("\nBắt đầu tạo bộ ba từ dữ liệu trích xuất...")

#feature_keys = ['English', 'Context_Objectives', 'Summarize_the_contents', 'Expected features']
# feature_keys_for_triplet = ['English', 'Context_Objectives', 'Summarize_the_contents', 'Expected features']
# Lặp qua từng dự án và tạo bộ ba
# data = {
        #     'English': '',
        #     'Context_Objectives': '',
        #     'Summarize_the_contents':'',
        #     'Expected features': '',     
        #     'filename': '',     
        # }
        
   
#lay data technology
# data_technology = []
# for project in all_projects_data:
#     technology_text = project.get('Technology/algorithm', '')
#     if technology_text:
#         data_technology.append({
#             'filename': project['filename'],
#             'technology': technology_text
#         })
# df_technology = pd.DataFrame(data_technology)
# df_technology.to_csv('D:/Traning AI/technology.csv', index=False, encoding='utf-8-sig')
# #trich xuat da ta        
# for i, anchor_project in enumerate(all_projects_data):
    
#     for anchor_key in feature_keys_for_triplet:
#         anchor_text = anchor_project.get(anchor_key, '')
#         if not anchor_text:
#             continue
#         # Lặp qua các feature *khác* của *cùng* anchor_project để chọn làm positive_text
#         for positive_key in feature_keys_for_triplet:
#             if positive_key == anchor_key:
#                 continue
#             positive_text = anchor_project.get(positive_key, '')
#             if not positive_text:
#                 continue
#             # Lặp qua từng đề tài *khác* để làm negative
#             other_project_indices = [idx for idx in range(len(all_projects_data)) if idx != i]
#             if not other_project_indices:
#                 continue
#             num_negatives_project = min(3, len(other_project_indices))  # Giới hạn số lượng negative dự án
#             selected_negative_indices = random.sample(other_project_indices, num_negatives_project)
#             for negative_index in selected_negative_indices:
#                 negative_project_data = all_projects_data[negative_index]
#                 # Negative_text là feature CÙNG LOẠI với anchor_text, nhưng từ đề tài khác
#                 negative_text = negative_project_data.get(anchor_key, '')
#                 if not negative_text:
#                     continue
#                 triplet_rows.append({
#                     'filename_anchor_feature': anchor_project['filename'] + " " + anchor_key,
#                     'anchor': anchor_text,
#                     'positive': positive_text,  # Positive là chính nó
#                     'negative': negative_text,  # Negative là dự án khác 
#                     'filename_negative_feature': negative_project_data['filename'] + " " + anchor_key
#                 })
        
        
        
                    
# print(f"Đã tạo {len(triplet_rows)} bộ ba từ dữ liệu trích xuất.")
# # 4. GHI dữ liệu vào file CSV
# print(f"\nĐang ghi dữ liệu vào file CSV: {output_csv_path}...")
# df = pd.DataFrame(triplet_rows)
# df.to_csv(output_csv_path, index=False, encoding='utf-8-sig')
# print("Đã ghi dữ liệu vào file CSV thành công!")
### close


def create_triple_rows(context_feature_list, feature_key):
    """
    Tạo bộ ba từ danh sách các văn bản và khóa tính năng.
    Trả về danh sách các dictionary chứa bộ ba.
    """
    triplet_rows = []
    dem=1
    if not context_feature_list or not feature_key:
        print("Cảnh báo: Danh sách paraphrase (context_feature_list) trống.")
        return triplet_rows
    # paraphrases_lookup['filename']['feature_key'] = [list_of_paraphrases]
    for i, anchor_project in enumerate(all_projects_data):
        anchor_filename = anchor_project.get('filename', '')
        anchor_text_original = anchor_project.get(feature_key, '')
        
        if not anchor_filename:
            print(f"Cảnh báo: Không có file {anchor_filename}.")
            continue
        if not anchor_text_original:
            print(f"Cảnh báo: Không có văn bản gốc cho {anchor_filename} với khóa {feature_key}.")
            continue
        positive_paraphases = []
        #Lấy danh sách paraphrase cho anchor_filename và feature_key
        if anchor_filename == context_feature_list[i]['filename']:
            print(f"Đang xử lý file: {anchor_filename} với khóa {feature_key}...")
            positive_paraphases = context_feature_list[i].get('list_parapharase', [])
        if not positive_paraphases:
            print(f"Cảnh báo: Không có paraphrase cho {anchor_filename} với khóa {feature_key}.")
            continue
        
        for positive_text in positive_paraphases:
            other_project_indices = [idx for idx in range(len(all_projects_data)) if idx != i]
            if not other_project_indices:
                continue
            num_negatives_project = min(3, len(other_project_indices))  # Giới hạn số lượng negative dự án
            selected_negative_indices = random.sample(other_project_indices, num_negatives_project)
            for negative_index in selected_negative_indices:
                print("Số bộ ba đã tạo:", dem)
                dem+=1
                negative_project = all_projects_data[negative_index]
                negative_filename = negative_project.get('filename', '')
                negative_text = negative_project.get(feature_key, '')
                if not negative_text or not negative_filename:
                    continue
                triplet_rows.append({
                    'filename_anchor_feature': f"{anchor_filename} {feature_key}",
                    'anchor': anchor_text_original,
                    'positive': positive_text,  # Positive là chính nó
                    'negative': negative_text,  # Negative là dự án khác 
                    'filename_negative_feature': f"{negative_filename} {feature_key}"
                })
    if not triplet_rows:
        print(f"Cảnh báo: Không tạo được bộ ba nào cho khóa {feature_key}.")
    print("Length of triplet_rows:", len(triplet_rows))
    return triplet_rows            
    

    

from AutoGetParaPhase import paraphrase_on_zerogpt
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.service import Service as ChromeService
from webdriver_manager.chrome import ChromeDriverManager
# 3. PARAPHRASE TỪNG DÒNG VĂN BẢN TRONG DANH SÁCH
#Khởi tạo danh sách gốc

# for project in all_projects_data:
#         if (len(project['Expected features'])> 300):
#             print(f"File {project['filename']} có nội dung tóm tắt quá dài ({len(project['Expected features'])} ký tự).")

  
# if __name__ == "__main__":
#     sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
    
#     triplet_rows = []
    
#     #trích xuất dữ liệu từ các file docx
#     folder_path = "D:/Traning AI/OK"
#     output_csv_path = "D:/Traning AI/output_triplet_in_progress_SP25.csv"
#     # 2. TRÍCH XUẤT DỮ LIỆU TỪ TẤT CẢ CÁC FILE DOCX
#     all_projects_data = []
#     print(f"Bắt đầu quét thư mục: {folder_path}")
#     for filename in os.listdir(folder_path):
#         if filename.endswith(".docx"):
#             file_path = os.path.join(folder_path, filename)
#             print(f"Đang xử lý file: {filename}...")
#             project_data = extract_project_info(file_path)
#             project_data['filename'] = filename  # Lưu tên file vào dữ liệu
#             if project_data:
#                 all_projects_data.append(project_data)
#                 print(f"Đã trích xuất dữ liệu từ file: {filename}")

#     print(f"\nĐã trích xuất thành công dữ liệu từ {len(all_projects_data)} đề tài.")
    
#     # In ra kết quả cuối cùng
#     context_Expected_features_original_list = []
#     for project in (all_projects_data):
#         new_data_point = {
#             'filename': project.get('filename', ''),
#             'text': project.get('Expected features', '')       
#         }
#         context_Expected_features_original_list.append(new_data_point)
    
#     # --- QUẢN LÝ TRÌNH DUYỆT ---
#     #Khởi tạo driver một lần duy nhất
#     print("Đang khởi tạo trình duyệt...")
#     driver = None # Khai báo trước để khối finally có thể truy cập
#     try:
#         options = webdriver.ChromeOptions()
#         options.add_argument("--start-maximized")
#         options.add_argument("--disable-blink-features=AutomationControlled")
#         options.add_experimental_option("excludeSwitches", ["enable-automation"])
#         options.add_experimental_option('useAutomationExtension', False)
#         driver = webdriver.Chrome(service=ChromeService(ChromeDriverManager().install()), options=options)
#         driver.execute_script("Object.defineProperty(navigator, 'webdriver', {get: () => undefined})")

#         # Mở trang web một lần
#         TARGET_URL = "https://www.zerogpt.com/paraphraser" # <- Thay URL đúng vào đây
#         driver.get(TARGET_URL)
#         print(f"✅ Đã truy cập thành công trang: {TARGET_URL}")
#         # ... (Xử lý cookie nếu cần) ...

#     except Exception as e:
#         print(f"❌ Lỗi nghiêm trọng khi khởi tạo trình duyệt: {e}")
#         # Thoát nếu không khởi tạo được trình duyệt
#         if driver:
#             driver.quit()
#         exit()
    
#     # Danh sách cuối cùng để lưu kết quả
#     context_Expected_features_new_list = []
    
#     # --- VÒNG LẶP XỬ LÝ CHÍNH THEO LOGIC MỚI ---
#     try:
#         # 1. Bắt đầu vòng lặp ngoài: duyệt qua từng file/văn bản gốc
#         for data_point in context_Expected_features_original_list:
#             filename = data_point.get('filename')
#             original_text = data_point.get('text')

#             if not original_text:
#                 print(f"⚠️ Bỏ qua file {filename} vì không có nội dung.")
#                 final_data_structure = {
#                     'filename': filename,
#                     'list_parapharase': []
#                 }
#                 context_Expected_features_new_list.append(final_data_structure)
#                 continue
            
#             print(f"\n--- Đang xử lý file: {filename} ---")
            
#             # 2. Tạo một danh sách tạm để lưu các phiên bản paraphrase của CHỈ file này
#             list_parapharase = []
            
#             # 3. Bắt đầu vòng lặp trong: paraphrase 3 lần
#             for i in range(3):
#                 print(f"  -> Đang paraphrase lần {i+1}/3...")
                
#                 # Gọi hàm paraphrase và truyền driver cùng văn bản gốc vào
#                 new_text = paraphrase_on_zerogpt(driver, original_text)
                
#                 # Kiểm tra kết quả trả về có hợp lệ không
#                 if new_text and new_text.strip():
#                     print(f"  -> Thành công!")
#                     list_parapharase.append(new_text)
#                 else:
#                     print(f"  -> ❌ Thất bại hoặc kết quả trống.")
            
#             # 4. Sau khi lặp xong, kiểm tra xem có thu được kết quả nào không
#             if list_parapharase:
#                 # 5. Tạo cấu trúc dictionary cuối cùng như bạn yêu cầu
#                 final_data_structure = {
#                     'filename': filename,
#                     'list_parapharase': list_parapharase
#                 }
                
#                 # 6. Thêm dictionary hoàn chỉnh này vào danh sách kết quả cuối cùng
#                 context_Expected_features_new_list.append(final_data_structure)
#             else:
#                 print(f"❌ Không thu được kết quả paraphrase nào cho file: {filename}")

#     finally:
#         # Đóng trình duyệt một lần duy nhất khi tất cả đã xong
#         print("\n✅ Hoàn thành tất cả các tác vụ. Đang đóng trình duyệt.")
#         if driver:
#             driver.quit()
#         for data_point in context_Expected_features_new_list:
#             print(f"\n--- File: {data_point['filename']} ---")
#             paraphrases = data_point.get('list_parapharase', [])
#             if paraphrases:
#                 print(f"  Số lượng paraphrase: {len(paraphrases)}")
#             else:
#                 print("  Không có paraphrase nào.")
#     tripplet_rows_feature = []
#     # 7. Tạo bộ ba từ danh sách paraphrase
#     tripplet_rows_feature = create_triple_rows(context_Expected_features_new_list, 'Expected features')
#     print(f"\nĐã tạo {len(tripplet_rows_feature)} bộ ba từ danh sách paraphrase.")
#     # 4. GHI dữ liệu vào file CSV
#     output_csv_path = "D:/Traning AI/output_triplets_with_feature_Expected_features_SP25.csv"
#     print(f"\nĐang ghi dữ liệu vào file CSV: {output_csv_path}...")
#     df = pd.DataFrame(tripplet_rows_feature)
#     df.to_csv(output_csv_path, index=False, encoding='utf-8-sig')
#     print("Đã ghi dữ liệu vào file CSV thành công!")
    
 
    
            
